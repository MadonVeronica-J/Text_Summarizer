from django.shortcuts import render,redirect
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.models import User
from .models import Summary
from django.contrib import messages 
from django.http import JsonResponse
input_text = ""
summary = ""
def user_login(request):
    print(input_text)
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']
        user = authenticate(request, username=username, password=password)

        if user is not None:
            login(request, user)
            return redirect('summarizer_dashboard')
        else:
            messages.error(request, "Invalid username or password.")
            return redirect('user_login')

    return render(request, 'login.html')

 
from transformers import pipeline

def summarizer_dashboard(request): 
    return render(request, 'summarizer_dashboard.html')

from sumy.parsers.plaintext import PlaintextParser 
from sumy.nlp.tokenizers import Tokenizer 
from sumy.summarizers.text_rank import TextRankSummarizer
from transformers import T5Tokenizer, T5ForConditionalGeneration
model_name = "google/flan-t5-base"
tokenizer = T5Tokenizer.from_pretrained(model_name)
model = T5ForConditionalGeneration.from_pretrained(model_name)

def abstract_summary(text,max_words):
    print('v')
    input_text = "summarize: " + text
    print(input_text)
    input_ids = tokenizer.encode(input_text, return_tensors="pt", max_length=1000, truncation=True)
    print(input_ids)

    summary_ids = model.generate(input_ids,min_length=30, max_length=max_words, num_beams=5, early_stopping=True)
    print(summary_ids)
    return tokenizer.decode(summary_ids[0], skip_special_tokens=True)

def extract_summary(text, num_sentences=3): 
    parser = PlaintextParser.from_string(text, Tokenizer("english")) 
    summarizer = TextRankSummarizer() 
    summary = summarizer(parser.document, num_sentences) 
    return " ".join(str(sentence) for sentence in summary)



def summarize_text(request): 
    if request.method == "POST": 
        text = request.POST.get("text") 
        summary_type = request.POST.get("summary_type") 
        value = int(request.POST.get("value"))
        
        if summary_type == "extract":
            summary_text = extract_summary(text, value)
            print(summary_text)
        else:
            summary_text = abstract_summary(text, value)
        
        summary_obj = Summary.objects.create(
            user=request.user,
            input_text = text,
            summary_text=summary_text,
            summary_type=summary_type
        )

        return JsonResponse({"summary": summary_text, "summary_id":summary_obj.id})

    return JsonResponse({"error": "Invalid request"})

from django.contrib.auth.hashers import make_password

def register(request):
    if request.method == "POST":
        username = request.POST["username"]
        password1 = request.POST["password1"]
        password2 = request.POST["password2"]

        if password1 != password2:
            messages.error(request, "Passwords do not match.")
            return redirect("register")

        if User.objects.filter(username=username).exists():
            messages.error(request, "Username already taken.")
            return redirect("register")


        user = User.objects.create(username=username,  password=make_password(password1))
        login(request, user)  # Auto-login after registration
        return redirect("user_login")  # Redirect to home page

    return render(request, "register.html")

def history_view(request):
    summaries = Summary.objects.filter(user=request.user).order_by('-created_at')
    return render(request, "history.html", {"summaries": summaries})

def user_logout(request):
    logout(request)
    return redirect('user_login')

from django.http import HttpResponse
from docx import Document
from fpdf import FPDF
from io import BytesIO

def download_word(request, summary_id):
    summary = Summary.objects.get(id=summary_id, user=request.user)
    print(summary.summary_text)
    doc = Document()
    doc.add_heading('Generated Summary', level=1)
    doc.add_paragraph(summary.summary_text)

    buffer = BytesIO()

    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(buffer.getvalue(),content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=summary_{summary_id}.docx'

   
    
    
    return response

def download_pdf(request, summary_id):
    summary = Summary.objects.get(id=summary_id, user=request.user)
    print(summary.summary_text)

    p = FPDF()
    p.add_page()
    p.set_font("Arial", size=12)
    p.multi_cell(0,10,summary.summary_text)

    
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename=summary_{summary_id}.pdf'
    p.output(dest='S').encode('latin1')
    response.write(p.output(dest='S').encode('latin1'))
    return response